import os
from docx import Document
from pypdf import PdfReader

docs_dir = os.path.dirname(os.path.abspath(__file__))

def extract_docx(filename, output_name):
    path = os.path.join(docs_dir, filename)
    out_path = os.path.join(docs_dir, output_name)
    if not os.path.exists(path):
        print(f"File {filename} not found.")
        return
    doc = Document(path)
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            text.append(" | ".join(row_data))
            
    with open(out_path, "w", encoding="utf-8") as f:
        f.write("\n".join(text))
    print(f"Extracted {filename} to {output_name}")

def extract_pdf(filename, output_name):
    path = os.path.join(docs_dir, filename)
    out_path = os.path.join(docs_dir, output_name)
    if not os.path.exists(path):
        print(f"File {filename} not found.")
        return
    reader = PdfReader(path)
    text = []
    for i, page in enumerate(reader.pages):
        text.append(f"--- Page {i+1} ---")
        text.append(page.extract_text() or "")
    
    with open(out_path, "w", encoding="utf-8") as f:
        f.write("\n".join(text))
    print(f"Extracted {filename} to {output_name}")

if __name__ == "__main__":
    extract_pdf("User_Story_Map_v1.2.pdf", "User_Story_Map_v1.2.md")
    extract_docx("Requisitos_funcionales_v1.2 (1).docx", "Requisitos_funcionales_v1.2.md")
    extract_docx("Nuclear-Patrones-dis (1).docx", "Nuclear-Patrones-dis.md")
